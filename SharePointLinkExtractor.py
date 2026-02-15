# https://github.com/gogeekery/Python-Scripts
# Crawls documents in a SharePoint site to find all links

# pip install msal requests python-docx PyPDF2 pandas cryptography

import io
import os
import re
import json
import time
import zipfile
import threading
import requests
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from urllib.parse import urlparse
from pathlib import Path
import msal
from docx import Document
import PyPDF2
import pandas as pd
from openpyxl import Workbook
import xml.etree.ElementTree as ET
from cryptography.fernet import Fernet

KEY_PATH = Path.home() / "sharepoint_link_extractor_config.key"
CONFIG_PATH = Path.home() / "sharepoint_link_extractor_config.enc"

DEFAULT_SITE = ""


def _generate_key():
    key = Fernet.generate_key()
    with open(KEY_PATH, "wb") as f:
        f.write(key)
    return key

def _load_key():
    if not KEY_PATH.exists():
        return _generate_key()
    with open(KEY_PATH, "rb") as f:
        return f.read()

fernet = Fernet(_load_key())



# ---------------- Utilities ----------------
def dedupe_display_text(text: str) -> str:
    """If text is a repetition of a smaller substring, return the substring."""
    n = len(text)
    if n == 0:
        return text
    pi = [0] * n
    for i in range(1, n):
        j = pi[i - 1]
        while j > 0 and text[i] != text[j]:
            j = pi[j - 1]
        if text[i] == text[j]:
            j += 1
        pi[i] = j
    pattern_len = n - pi[-1]
    if pattern_len < n and n % pattern_len == 0:
        return text[:pattern_len]
    return text

# ---------------- GUI / App ----------------
class LinkExtractorUI:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("SharePoint Link Extractor")
        self.root.geometry("980x760")

        self.stop_flag = False
        self.total_files_scanned = 0
        self.total_links_found = 0
        self.links = []  # dicts: site, file_name, file_web_url, found_url, display_text

        # Regex: http(s) and www. patterns
        self.url_pattern = re.compile(r'(https?://[^\s<>")]+|www\.[^\s<>")]+)', re.IGNORECASE)

        self._build_ui()
        self._load_config()

    # ---------------- UI ----------------
    def _build_ui(self):
        header = ttk.Label(self.root, text="SharePoint Link Extractor", font=("Segoe UI", 20, "bold"))
        header.pack(pady=12)

        conn_frame = ttk.LabelFrame(self.root, text="Connection / App Credentials")
        conn_frame.pack(fill="x", padx=20, pady=8)

        ttk.Label(conn_frame, text="Site URL:").grid(row=0, column=0, sticky="w", padx=8, pady=6)
        self.site_entry = ttk.Entry(conn_frame, width=80)
        self.site_entry.insert(0, DEFAULT_SITE)
        self.site_entry.grid(row=0, column=1, padx=8, pady=6, columnspan=3)

        ttk.Label(conn_frame, text="Tenant ID:").grid(row=1, column=0, sticky="w", padx=8, pady=6)
        self.tenant_entry = ttk.Entry(conn_frame, width=36)
        self.tenant_entry.grid(row=1, column=1, padx=8, pady=6)

        ttk.Label(conn_frame, text="Client ID:").grid(row=1, column=2, sticky="w", padx=8, pady=6)
        self.client_entry = ttk.Entry(conn_frame, width=36)
        self.client_entry.grid(row=1, column=3, padx=8, pady=6)

        ttk.Label(conn_frame, text="Client Secret:").grid(row=2, column=0, sticky="w", padx=8, pady=6)
        self.secret_entry = ttk.Entry(conn_frame, width=80, show="*")
        self.secret_entry.grid(row=2, column=1, padx=8, pady=6, columnspan=3)

        self.save_config_btn = ttk.Button(conn_frame, text="Save Config", command=self._save_config)
        self.save_config_btn.grid(row=3, column=1, sticky="w", padx=8, pady=6)
        self.clear_config_btn = ttk.Button(conn_frame, text="Clear Saved Config", command=self._clear_config)
        self.clear_config_btn.grid(row=3, column=2, sticky="w", padx=8, pady=6)

        toggles_frame = ttk.Frame(self.root)
        toggles_frame.pack(fill="x", padx=20, pady=6)
        self.scan_docx_var = tk.BooleanVar(value=True)
        self.scan_pdf_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(toggles_frame, text="Scan DOCX", variable=self.scan_docx_var).pack(side="left", padx=6)
        ttk.Checkbutton(toggles_frame, text="Scan PDF", variable=self.scan_pdf_var).pack(side="left", padx=6)

        control_frame = ttk.Frame(self.root)
        control_frame.pack(fill="x", padx=20, pady=8)
        self.start_button = ttk.Button(control_frame, text="Start Scan", command=self.start_scan)
        self.start_button.pack(side="left", padx=6)
        self.stop_button = ttk.Button(control_frame, text="Stop", command=self.stop_scan, state="disabled")
        self.stop_button.pack(side="left", padx=6)
        self.export_csv_btn = ttk.Button(control_frame, text="Export CSV", command=self.export_csv, state="disabled")
        self.export_csv_btn.pack(side="left", padx=6)
        self.export_xlsx_btn = ttk.Button(control_frame, text="Export XLSX", command=self.export_xlsx, state="disabled")
        self.export_xlsx_btn.pack(side="left", padx=6)
        self.status_label = ttk.Label(control_frame, text="Status: Idle", foreground="blue")
        self.status_label.pack(side="right")

        progress_frame = ttk.LabelFrame(self.root, text="Progress")
        progress_frame.pack(fill="x", padx=20, pady=8)
        self.progress = ttk.Progressbar(progress_frame, orient="horizontal", mode="determinate")
        self.progress.pack(fill="x", padx=10, pady=6)
        self.progress_label = ttk.Label(progress_frame, text="Files Scanned: 0 | Links Found: 0 | Total Files: 0")
        self.progress_label.pack(padx=10, pady=6)

        log_frame = ttk.LabelFrame(self.root, text="Activity Log")
        log_frame.pack(fill="both", expand=True, padx=20, pady=8)
        self.log_text = tk.Text(log_frame, wrap="none")
        self.log_text.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(log_frame, command=self.log_text.yview)
        scrollbar.pack(side="right", fill="y")
        self.log_text.configure(yscrollcommand=scrollbar.set)

    def run(self):
        self.root.mainloop()

    # ---------------- Config ----------------
    def _save_config(self):
        cfg = {
            "site_url": self.site_entry.get().strip(),
            "tenant_id": self.tenant_entry.get().strip(),
            "client_id": self.client_entry.get().strip(),
            "client_secret": self.secret_entry.get().strip()
        }
        try:
            data = json.dumps(cfg).encode("utf-8")       # Convert dict to bytes
            encrypted = fernet.encrypt(data)             # Encrypt bytes
            with open(CONFIG_PATH, "wb") as f:          # Write encrypted bytes
                f.write(encrypted)
            self._log(f"Config saved (encrypted) to {CONFIG_PATH}\n")
            messagebox.showinfo("Saved", "Configuration saved securely.")
        except Exception as e:
            self._log(f"Failed to save config: {e}\n")
            messagebox.showerror("Error", f"Failed to save config: {e}")

    def _load_config(self):
        try:
            if not CONFIG_PATH.exists():
                self._log(f"No saved config found\n")
                return None
            with open(CONFIG_PATH, "rb") as f:
                encrypted = f.read()
            decrypted = fernet.decrypt(encrypted)
            cfg = json.loads(decrypted.decode("utf-8"))

            # Populate the entry fields with decrypted values
            self.site_entry.delete(0, tk.END)
            self.site_entry.insert(0, cfg.get("site_url", ""))

            self.tenant_entry.delete(0, tk.END)
            self.tenant_entry.insert(0, cfg.get("tenant_id", ""))

            self.client_entry.delete(0, tk.END)
            self.client_entry.insert(0, cfg.get("client_id", ""))

            self.secret_entry.delete(0, tk.END)
            self.secret_entry.insert(0, cfg.get("client_secret", ""))

            self._log(f"Config loaded successfully from {CONFIG_PATH}\n")
            return cfg
        except Exception as e:
            self._log(f"Failed to load config: {e}\n")
            messagebox.showerror("Error", f"Failed to load config: {e}")
            return None


    def _clear_config(self):
        try:
            if CONFIG_PATH.exists():
                CONFIG_PATH.unlink()
            if KEY_PATH.exists():
                KEY_PATH.unlink()
            self._log("Encrypted config cleared.\n")
            messagebox.showinfo("Cleared", "Saved configuration cleared.")
        except Exception as e:
            self._log(f"Failed to clear config: {e}\n")
            messagebox.showerror("Error", f"Failed to clear config: {e}")

    # ---------------- Start / Stop ----------------
    def start_scan(self):
        site = self.site_entry.get().strip()
        tenant = self.tenant_entry.get().strip()
        client_id = self.client_entry.get().strip()
        client_secret = self.secret_entry.get().strip()
        scan_docx = self.scan_docx_var.get()
        scan_pdf = self.scan_pdf_var.get()

        if not site or not tenant or not client_id or not client_secret:
            messagebox.showerror("Missing Info", "Please provide Site URL, Tenant ID, Client ID, and Client Secret.")
            return
        if not (scan_docx or scan_pdf):
            messagebox.showerror("No File Types", "Enable at least one file type to scan (DOCX or PDF).")
            return

        self.stop_flag = False
        self.total_files_scanned = 0
        self.total_links_found = 0
        self.links = []
        self.start_button.config(state="disabled")
        self.stop_button.config(state="normal")
        self.export_csv_btn.config(state="disabled")
        self.export_xlsx_btn.config(state="disabled")
        self._update_status("Preparing scan…")

        threading.Thread(target=self._scan_worker, args=(site, tenant, client_id, client_secret, scan_docx, scan_pdf), daemon=True).start()

    def stop_scan(self):
        self.stop_flag = True
        self._log("Stop requested...\n")
        self._update_status("Stopping...")

    # ---------------- Logging / UI helpers ----------------
    def _log(self, text):
        timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
        self.log_text.insert("end", f"[{timestamp}] {text}")
        self.log_text.see("end")

    def _update_progress(self, total_files=0):
        self.progress_label.config(text=f"Files Scanned: {self.total_files_scanned} | Links Found: {self.total_links_found} | Total Files: {total_files}")
        try:
            self.progress['value'] = self.total_files_scanned
        except Exception:
            pass

    def _update_status(self, text):
        self.status_label.config(text=f"Status: {text}")

    # ---------------- Graph / MSAL helpers ----------------
    def _acquire_token(self, tenant_id, client_id, client_secret, scope=["https://graph.microsoft.com/.default"]):
        authority = f"https://login.microsoftonline.com/{tenant_id}"
        app = msal.ConfidentialClientApplication(client_id, authority=authority, client_credential=client_secret)
        result = app.acquire_token_for_client(scopes=scope)
        if "access_token" in result:
            return result["access_token"]
        else:
            raise Exception(f"Token acquisition failed: {result.get('error_description') or result}")

    def _get_site_resource(self, access_token, site_url):
        parsed = urlparse(site_url)
        hostname = parsed.netloc
        path = parsed.path.rstrip('/')
        if not path:
            path = '/'
        graph_url = f"https://graph.microsoft.com/v1.0/sites/{hostname}:{path}"
        headers = {"Authorization": f"Bearer {access_token}"}
        r = requests.get(graph_url, headers=headers)
        if r.status_code == 200:
            return r.json()
        else:
            raise Exception(f"Failed to resolve site: {r.status_code} {r.text}")

    def _list_drive_children(self, access_token, drive_id, item_id=None):
        headers = {"Authorization": f"Bearer {access_token}"}
        if item_id:
            url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/children"
        else:
            url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root/children"
        while url:
            r = requests.get(url, headers=headers)
            if r.status_code != 200:
                raise Exception(f"List children failed: {r.status_code} {r.text}")
            data = r.json()
            for item in data.get("value", []):
                yield item
            url = data.get("@odata.nextLink")

    def _download_drive_item(self, access_token, drive_id, item_id):
        headers = {"Authorization": f"Bearer {access_token}"}
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/content"
        r = requests.get(url, headers=headers, stream=True)
        if r.status_code in (200, 302):
            return r.content
        else:
            raise Exception(f"Download failed: {r.status_code} {r.text}")

    # ---------------- Count target files ----------------
    def _count_target_files(self, access_token, drive_id, scan_docx, scan_pdf):
        count = 0
        stack = [None]
        while stack and not self.stop_flag:
            current = stack.pop()
            for item in self._list_drive_children(access_token, drive_id, current):
                if self.stop_flag:
                    break
                item_name = item.get("name", "").lower()
                item_id = item.get("id")
                if item.get("folder"):
                    stack.append(item_id)
                elif item.get("file"):
                    if (scan_docx and item_name.endswith(".docx")) or (scan_pdf and item_name.endswith(".pdf")):
                        count += 1
        return count

    # ---------------- Extraction (bytes-based) ----------------
    def _extract_links_from_docx_bytes(self, b: bytes):
        links = []
        # Try python-docx first
        try:
            doc = Document(io.BytesIO(b))
            # paragraphs and hyperlink rels
            for para in doc.paragraphs:
                # hyperlink elements inside paragraph
                try:
                    hyperlink_elements = para._element.xpath('.//w:hyperlink')
                except Exception:
                    hyperlink_elements = []
                for hl in hyperlink_elements:
                    rid = hl.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    url = ""
                    if rid and (rid in doc.part.rels):
                        url = doc.part.rels[rid].target_ref
                    display_text = "".join(hl.itertext())
                    links.append((url or "", display_text or ""))
                # plain-text URLs in paragraph
                for m in re.finditer(self.url_pattern, para.text or ""):
                    found_url = m.group()
                    if not any(found_url == u for u, _ in links):
                        links.append((found_url, found_url))
            # tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            try:
                                hyperlink_elements = para._element.xpath('.//w:hyperlink')
                            except Exception:
                                hyperlink_elements = []
                            for hl in hyperlink_elements:
                                rid = hl.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                                url = ""
                                if rid and (rid in doc.part.rels):
                                    url = doc.part.rels[rid].target_ref
                                display_text = "".join(hl.itertext())
                                links.append((url or "", display_text or ""))
                            for m in re.finditer(self.url_pattern, para.text or ""):
                                found_url = m.group()
                                if not any(found_url == u for u, _ in links):
                                    links.append((found_url, found_url))
            # also include any rels not captured above
            try:
                rels = doc.part.rels
                for rel in rels.values():
                    if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink":
                        target = rel.target_ref
                        if target and not any(target == u for u, _ in links):
                            links.append((target, target))
            except Exception:
                pass
        except Exception:
            # fallback to XML parsing from bytes
            links.extend(self._extract_docx_links_xml_bytes(b))
        # normalize and dedupe by URL, keep first display_text
        seen = {}
        for url, disp in links:
            if not url and disp:
                # if URL empty but display_text contains a URL, extract it
                m = re.search(self.url_pattern, disp or "")
                if m:
                    url = m.group()
            if url:
                if url not in seen:
                    seen[url] = disp
        return [(u, seen[u]) for u in seen]

    def _extract_docx_links_xml_bytes(self, b: bytes):
        links = []
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        try:
            with zipfile.ZipFile(io.BytesIO(b)) as z:
                xml_content = z.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                # load relationships
                rels = {}
                try:
                    rel_xml = z.read("word/_rels/document.xml.rels")
                    rel_tree = ET.fromstring(rel_xml)
                    for rel in rel_tree.findall(".//"):
                        rId = rel.get("Id") or rel.get("Id".lower())
                        target = rel.get("Target")
                        if rId and target:
                            rels[rId] = target
                except Exception:
                    pass
                # find hyperlink elements
                for hl in tree.findall('.//w:hyperlink', namespaces):
                    rId = hl.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    url = rels.get(rId, "") if rId else ""
                    display_text = "".join(hl.itertext())
                    links.append((url, display_text))
                # regex on full text
                full_text = "".join(tree.itertext())
                for m in re.finditer(self.url_pattern, full_text):
                    found_url = m.group()
                    if not any(found_url == u for u, _ in links):
                        links.append((found_url, found_url))
        except Exception:
            pass
        # normalize/dedupe
        seen = {}
        for url, disp in links:
            if url and url not in seen:
                seen[url] = disp
        return [(u, seen[u]) for u in seen]

    def _extract_links_from_pdf_bytes(self, b: bytes):
        links = []
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(b))
            # text extraction
            for page in reader.pages:
                try:
                    text = page.extract_text() or ""
                    for m in re.finditer(self.url_pattern, text):
                        found_url = m.group()
                        if not any(found_url == u for u, _ in links):
                            links.append((found_url, found_url))
                except Exception:
                    continue
            # annotations (URIs and contents)
            try:
                for page in reader.pages:
                    annots = page.get("/Annots")
                    if not annots:
                        continue
                    for a in annots:
                        obj = a.get_object()
                        action = obj.get("/A")
                        if isinstance(action, PyPDF2.generic.DictionaryObject):
                            uri = action.get("/URI")
                            if uri:
                                display_text = obj.get("/Contents", "") or uri
                                if not any(uri == u for u, _ in links):
                                    links.append((uri, display_text))
            except Exception:
                pass
        except Exception:
            # fallback: regex on raw bytes
            try:
                text = b.decode('utf-8', errors='ignore')
                for m in re.finditer(self.url_pattern, text):
                    found_url = m.group()
                    if not any(found_url == u for u, _ in links):
                        links.append((found_url, found_url))
            except Exception:
                pass
        # normalize/dedupe
        seen = {}
        for url, disp in links:
            if url and url not in seen:
                seen[url] = disp
        return [(u, seen[u]) for u in seen]

    # ---------------- Crawl worker ----------------
    def _scan_worker(self, site_url, tenant_id, client_id, client_secret, scan_docx, scan_pdf):
        total_files = 0
        try:
            token = self._acquire_token(tenant_id, client_id, client_secret)
            self._update_status("Token acquired")
            self._log("Token acquired.\n")

            site = self._get_site_resource(token, site_url)
            site_id = site.get("id")
            self._log(f"Resolved site id: {site_id}\n")

            headers = {"Authorization": f"Bearer {token}"}
            drive_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive"
            r = requests.get(drive_url, headers=headers)
            if r.status_code != 200:
                raise Exception(f"Failed to get drive: {r.status_code} {r.text}")
            drive = r.json()
            drive_id = drive.get("id")
            self._log(f"Using drive id: {drive_id}\n")

            # count files
            self._update_status("Counting target files...")
            self._log("Counting files to scan (this may take a moment)...\n")
            total_files = self._count_target_files(token, drive_id, scan_docx, scan_pdf)
            self._log(f"Total target files: {total_files}\n")
            self.progress['maximum'] = max(1, total_files)
            self._update_progress(total_files=total_files)

            # traverse and download
            self._update_status(f"Scanning (0/{total_files})")
            stack = [None]
            while stack and not self.stop_flag:
                current = stack.pop()
                for item in self._list_drive_children(token, drive_id, current):
                    if self.stop_flag:
                        break
                    item_name = item.get("name")
                    item_id = item.get("id")
                    item_folder = item.get("folder")
                    item_file = item.get("file")
                    web_url = item.get("webUrl", "")
                    if item_folder:
                        stack.append(item_id)
                    elif item_file:
                        lower = (item_name or "").lower()
                        is_docx = lower.endswith(".docx")
                        is_pdf = lower.endswith(".pdf")
                        if (is_docx and scan_docx) or (is_pdf and scan_pdf):
                            self._log(f"Downloading: {web_url}\n")
                            try:
                                content = self._download_drive_item(token, drive_id, item_id)
                                self.total_files_scanned += 1
                                found = []
                                if is_docx:
                                    found = self._extract_links_from_docx_bytes(content)
                                elif is_pdf:
                                    found = self._extract_links_from_pdf_bytes(content)

                                for u, disp in found:
                                    self.links.append({
                                        "site": site_url,
                                        "file_name": item_name,
                                        "file_web_url": web_url,
                                        "found_url": u,
                                        "display_text": dedupe_display_text(disp or "")
                                    })
                                self.total_links_found += len(found)
                                self._log(f"Scanned {item_name}: {len(found)} links found\n")
                                self._update_progress(total_files=total_files)
                                self._update_status(f"Scanning ({self.total_files_scanned}/{total_files})")
                            except Exception as ex:
                                self._log(f"Error downloading/parsing {item_name}: {ex}\n")
                # end for children
            # end while

            if self.stop_flag:
                self._update_status("Stopped by user")
                self._log("Scan stopped by user.\n")
            else:
                self._update_status("Completed")
                self._log("Scan completed.\n")

        except Exception as e:
            self._log(f"Fatal error: {e}\n")
            self._update_status("Error")
        finally:
            self.start_button.config(state="normal")
            self.stop_button.config(state="disabled")
            if self.links:
                self.export_csv_btn.config(state="normal")
                self.export_xlsx_btn.config(state="normal")
            self._update_progress(total_files=total_files)

    # ---------------- Export ----------------
    def export_csv(self):
        if not self.links:
            messagebox.showinfo("No Data", "No links to export.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV files","*.csv")])
        if not path:
            return
        df = pd.DataFrame(self.links)
        df.to_csv(path, index=False)
        self._log(f"Exported {len(self.links)} rows to {path}\n")
        messagebox.showinfo("Exported", f"Exported {len(self.links)} rows to {path}")

    def export_xlsx(self):
        if not self.links:
            messagebox.showinfo("No Data", "No links to export.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Workbook","*.xlsx")])
        if not path:
            return
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Links"
            ws.append(["Site", "Source File", "File URL", "Found URL", "Display Text"])
            for row in self.links:
                ws.append([row.get("site"), row.get("file_name"), row.get("file_web_url"), row.get("found_url"), row.get("display_text")])
            wb.save(path)
            self._log(f"Exported {len(self.links)} rows to {path}\n")
            messagebox.showinfo("Exported", f"Exported {len(self.links)} rows to {path}")
        except Exception as e:
            self._log(f"Failed to export XLSX: {e}\n")
            messagebox.showerror("Export Error", f"Failed to export XLSX: {e}")

# ---------------- Run ----------------
if __name__ == "__main__":
    app = LinkExtractorUI()
    app.run()
